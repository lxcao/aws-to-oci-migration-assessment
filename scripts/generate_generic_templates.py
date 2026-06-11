#!/usr/bin/env python3
"""Generate generic AWS-to-OCI migration survey, assessment, and SOW templates."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


HEADERS = [
    "ID / 编号",
    "Dimension / 调研维度",
    "Question / 调研问题",
    "Customer Input Required / 客户需提供的信息",
    "AWS Current State / AWS现状",
    "OCI Target / OCI目标",
    "Assessment Focus / 评估重点",
    "Feedback / 反馈",
    "Initial Risk / 初步风险",
    "Notes / 备注",
]


SHEETS = {
    "App Architecture": [
        ("Business criticality / 业务重要性", "What is the business system, owner, environment, and criticality?", "System list, owner, environment, SLA/RTO/RPO"),
        ("Dependencies / 依赖关系", "What upstream/downstream systems does this workload depend on?", "Dependency diagram, protocol, endpoint list"),
        ("Release process / 发布流程", "How is the workload deployed and rolled back today?", "CI/CD flow, rollback procedure"),
    ],
    "Networking": [
        ("VPC design / VPC设计", "What are the VPC, subnet, route, NAT, gateway, and peering designs?", "CIDR, route tables, subnet list, connectivity diagram"),
        ("Security controls / 安全控制", "How are security groups, NACLs, firewalls, and east-west controls configured?", "Rule export, ports, source/destination matrix"),
        ("Hybrid connectivity / 混合连接", "Are VPN, Direct Connect, DNS forwarding, or on-premises routes required?", "Connectivity diagram, DNS rules, routing requirements"),
    ],
    "Load Balancing": [
        ("Listener and routing / 监听与路由", "Which ALB/NLB/CLB listeners, protocols, host/path rules, and certificates are used?", "Listener export, certificate list, routing rules"),
        ("Targets / 后端目标", "How are backend targets registered and health-checked?", "Target groups, health check settings, backend ports"),
        ("Traffic profile / 流量画像", "What are peak connections, throughput, idle timeout, and logging requirements?", "CloudWatch metrics, access log sample"),
    ],
    "ECS Containers": [
        ("Cluster and services / 集群与服务", "What ECS clusters, services, task definitions, launch types, and autoscaling policies are used?", "Cluster/service export, task definitions"),
        ("Images and config / 镜像与配置", "Where are images, secrets, environment variables, and runtime configs managed?", "Registry, secret store, env/config list"),
        ("Target platform / 目标平台", "Should the target be OKE, Container Instances, or Compute-based runtime?", "Operational preference, Kubernetes readiness, CI/CD constraints"),
    ],
    "S3 Object Storage": [
        ("Bucket inventory / 桶清单", "What buckets, object counts, sizes, storage classes, and growth rates exist?", "Bucket inventory, size report"),
        ("Access and policy / 访问与策略", "How are bucket policies, IAM, public access, pre-signed URLs, and encryption used?", "Policy export, KMS usage, access pattern"),
        ("Migration method / 迁移方式", "What consistency, validation, bandwidth, and cutover requirements apply?", "Data volume, change rate, cutover window"),
    ],
    "RDS MySQL": [
        ("Engine and config / 引擎与配置", "What MySQL version, parameter groups, charset, collation, and extensions are used?", "DB version, parameter export, schema inventory"),
        ("Performance / 性能", "What are database size, IOPS, CPU, memory, connections, and top SQL patterns?", "Performance metrics, slow SQL summary"),
        ("HA and backup / 高可用与备份", "What Multi-AZ, backup, PITR, replication, and maintenance requirements exist?", "Backup policy, HA config, maintenance window"),
    ],
    "Redis": [
        ("Topology / 拓扑", "What Redis version, cluster mode, shards, replicas, and node sizing are used?", "Cluster config, node sizing, memory usage"),
        ("Features / 功能特性", "Are persistence, pub/sub, streams, Lua scripts, TLS, AUTH, or eviction policies used?", "Parameter group, feature list, client usage"),
        ("Cutover / 割接", "How will endpoints, DNS, client retry, and cache warm-up be handled?", "Endpoint list, client settings, cutover plan"),
    ],
    "EFS File Storage": [
        ("Inventory / 清单", "What file systems, sizes, file counts, throughput, mount targets, and clients exist?", "File system inventory, client list"),
        ("Permissions / 权限", "What NFS, POSIX UID/GID, access point, and security controls are required?", "Mount options, permission model"),
        ("Data copy / 数据复制", "What migration copy, validation, and downtime approach is acceptable?", "Copy method, validation method, downtime window"),
    ],
    "Self-managed VM OS Storage": [
        ("Compute inventory / 计算清单", "What EC2 instances, OS versions, shapes, and attached volumes support self-managed components?", "Instance inventory, OS list, volume list"),
        ("OS and services / 操作系统与服务", "What system services, agents, cron jobs, ports, certificates, and startup scripts are required?", "Runbook, service list, port matrix"),
        ("Block storage / 块存储", "What volume type, size, IOPS, filesystem, snapshot, and backup requirements exist?", "Volume metrics, backup policy"),
    ],
    "NGINX to API Gateway": [
        ("Routes / 路由", "What NGINX host/path routes, rewrites, upstreams, and TLS settings are used?", "NGINX config, route inventory"),
        ("Auth and policy / 认证与策略", "What authentication, authorization, rate limit, CORS, and transform logic exists?", "Plugin/module list, policy config"),
        ("Fit-gap / 差异分析", "Which NGINX behaviors require OCI API Gateway mapping or custom implementation?", "Unsupported modules, custom scripts"),
    ],
    "Kafka to OCI Streaming": [
        ("Cluster and topics / 集群与主题", "What Kafka broker, topic, partition, replication, retention, and throughput profiles exist?", "Topic export, broker config, throughput metrics"),
        ("Clients / 客户端", "Which producers, consumers, consumer groups, connectors, and schema dependencies exist?", "Client list, connector list, schema registry usage"),
        ("Migration approach / 迁移方式", "Should migration use mirror, replay, dual-write, phased cutover, or self-managed target?", "RPO/RTO, lag tolerance, compatibility needs"),
    ],
    "Cutover Security Ops": [
        ("Cutover / 割接", "What migration waves, DNS changes, validation steps, rollback criteria, and signoffs are required?", "Wave plan, DNS TTL, validation checklist"),
        ("IAM and secrets / IAM与密钥", "What IAM roles, policies, KMS/secrets, certificates, and privileged access are required?", "IAM export, secret inventory, certificate list"),
        ("Monitoring and backup / 监控与备份", "What logs, metrics, alarms, dashboards, backup, restore, and DR tests are required?", "Monitoring inventory, backup policy"),
    ],
}


def build_workbook(path: Path) -> None:
    wb = Workbook()
    default = wb.active
    wb.remove(default)

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    risk_values = '"Low / 低,Medium / 中,High / 高,Critical / 关键"'

    for sheet_name, rows in SHEETS.items():
        ws = wb.create_sheet(sheet_name[:31])
        ws.append(HEADERS)
        for idx, (dimension, question, input_required) in enumerate(rows, start=1):
            ws.append([
                f"{sheet_name[:3].upper()}-{idx:02d}",
                dimension,
                question,
                input_required,
                "",
                "",
                "Assess compatibility, migration impact, risk, and required OCI design decisions.",
                "",
                "",
                "",
            ])
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        widths = [15, 26, 55, 42, 34, 34, 42, 42, 18, 30]
        for i, width in enumerate(widths, start=1):
            ws.column_dimensions[get_column_letter(i)].width = width
        for row_idx in range(2, ws.max_row + 1):
            ws.row_dimensions[row_idx].height = 54
        from openpyxl.worksheet.datavalidation import DataValidation

        dv = DataValidation(type="list", formula1=risk_values, allow_blank=True)
        ws.add_data_validation(dv)
        dv.add(f"I2:I{max(ws.max_row, 200)}")

    wb.save(path)


def setup_document_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    for style_name, size in [("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    doc.add_paragraph()


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for idx, (key, value) in enumerate(rows):
        table.cell(idx, 0).text = key
        table.cell(idx, 1).text = value
    doc.add_paragraph()


def build_report(path: Path) -> None:
    doc = Document()
    setup_document_styles(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    add_title(
        doc,
        "[Customer Name] AWS to OCI Migration Discovery and Feasibility Assessment",
        "[客户名称] AWS 到 OCI 迁移调研与可行性评估报告",
    )
    add_kv_table(doc, [
        ("Customer / 客户", "[Customer Name]"),
        ("Project / 项目", "[Project Name]"),
        ("Version / 版本", "V0.1"),
        ("Date / 日期", "[YYYY-MM-DD]"),
        ("Prepared by / 编写人", "[Team / Role]"),
    ])

    sections = [
        ("1. Executive Summary / 摘要", "Summarize migration feasibility, top risks, recommended migration path, and next actions."),
        ("2. Scope and Assumptions / 范围与假设", "List assessed AWS services, target OCI placeholders, assumptions, and exclusions."),
        ("3. Current AWS State Summary / AWS现状摘要", "Summarize customer feedback by service and identify missing information."),
        ("4. OCI Target Architecture Options / OCI目标架构方案", "Describe recommended OCI target services and alternatives where mapping is not one-to-one."),
        ("5. Service-by-Service Feasibility Assessment / 分服务可行性评估", "Use a table with AWS current state, OCI target, feasibility, risk, key issues, and mitigation."),
        ("6. Migration Approach and Tools / 迁移方案与工具", "Describe landing zone, network, workload, data, validation, cutover, and rollback approach."),
        ("7. Risk Register / 风险清单", "Track risk ID, risk description, severity, impact, mitigation, owner, and status."),
        ("8. POC and Acceptance Matrix / POC与验收矩阵", "Define validation objectives, success criteria, test method, and owner."),
        ("9. Migration Roadmap / 迁移路线图", "Group work into assessment, design, build, pilot, migration waves, stabilization, and handover."),
        ("10. Open Questions / 待确认事项", "Convert missing or unclear survey feedback into concrete customer questions."),
        ("11. References / 参考资料", "List official OCI documentation and customer-provided materials used for the assessment."),
    ]
    for heading, body in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
        if "table" in body.lower() or "Track risk" in body or "Define validation" in body:
            table = doc.add_table(rows=2, cols=4)
            table.style = "Table Grid"
            headers = ["Item / 项目", "Description / 描述", "Risk / 风险", "Notes / 备注"]
            for col, header in enumerate(headers):
                table.cell(0, col).text = header
            for col in range(4):
                table.cell(1, col).text = "[TBD]"
            doc.add_paragraph()
    doc.save(path)


def build_sow(path: Path) -> None:
    doc = Document()
    setup_document_styles(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    add_title(
        doc,
        "[Customer Name] Cloud Platform Migration Services Statement of Work",
        "[客户名称] 云平台迁移服务工作说明书",
    )
    add_kv_table(doc, [
        ("Customer / 客户", "[Customer Name]"),
        ("Project / 项目", "[Project Name]"),
        ("Version / 版本", "V0.1"),
        ("Date / 日期", "[YYYY-MM-DD]"),
        ("Based on / 依据", "[Migration Assessment Report Version]"),
    ])

    sections = [
        ("1. Introduction / 简介", "Define the purpose, document basis, and service relationship."),
        ("2. Project Background / 项目背景", "Describe the AWS-to-OCI migration context and business drivers."),
        ("3. Project Objectives / 项目目标", "Define target outcomes for OCI foundation, migration, validation, and handover."),
        ("4. Scope of Services / 服务范围", "List in-scope service areas, workstreams, and migration objects."),
        ("5. Out of Scope / 非服务范围", "List explicit exclusions such as app refactoring, long-term managed operations, or third-party license procurement unless contracted."),
        ("6. Delivery Methodology / 交付方法", "Describe discovery validation, design, build, pilot, migration waves, cutover, stabilization, and handover."),
        ("7. Deliverables / 交付物", "List deliverables such as design, runbook, migration plan, validation report, and handover material."),
        ("8. Roles and Responsibilities / 双方职责", "Define customer, provider, and third-party responsibilities with RACI if needed."),
        ("9. Timeline and Milestones / 项目计划与里程碑", "Provide milestone placeholders and dependency notes."),
        ("10. Acceptance Criteria / 验收标准", "State objective acceptance criteria for each deliverable and migrated workload."),
        ("11. Assumptions and Dependencies / 假设与依赖", "List customer access, information, downtime approval, OCI availability, and network prerequisites."),
        ("12. Change Management / 变更管理", "Define how scope, timeline, workload count, and acceptance criteria changes are controlled."),
        ("13. Risks and Mitigations / 风险与缓解", "Reference the migration assessment risk register and list SOW-level mitigations."),
        ("14. Confidentiality, Security, and Compliance / 保密、安全与合规", "Define credential handling, data protection, and compliance boundaries."),
        ("15. Exclusions and Commercial Notes / 除外事项与商务说明", "Keep commercial terms aligned with the approved contract process."),
    ]
    for heading, body in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
        if "List" in body or "Define" in body or "Provide" in body:
            table = doc.add_table(rows=2, cols=3)
            table.style = "Table Grid"
            headers = ["Area / 范围", "Responsibility or Criteria / 职责或标准", "Notes / 备注"]
            for col, header in enumerate(headers):
                table.cell(0, col).text = header
            for col in range(3):
                table.cell(1, col).text = "[TBD]"
            doc.add_paragraph()
    doc.save(path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-dir", help="Directory for generated template assets")
    args = parser.parse_args()

    default_output_dir = Path(__file__).resolve().parent.parent / "assets"
    output_dir = Path(args.output_dir).expanduser().resolve() if args.output_dir else default_output_dir
    output_dir.mkdir(parents=True, exist_ok=True)

    build_workbook(output_dir / "aws_to_oci_migration_survey_template.xlsx")
    build_report(output_dir / "aws_to_oci_feasibility_report_template.docx")
    build_sow(output_dir / "aws_to_oci_migration_sow_template.docx")

    print(f"Generated templates in {output_dir}")


if __name__ == "__main__":
    main()
